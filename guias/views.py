from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import JsonResponse, HttpResponse, Http404
from django.core.paginator import Paginator
from django.db.models import Q
from django.conf import settings
from django.utils import timezone
import os
import json
from datetime import datetime
import subprocess
import tempfile
import shutil

# Importaciones para generación de documentos
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor  # Agregar RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from htmldocx import HtmlToDocx  # Para convertir HTML de CKEditor a Word
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    DOCX_AVAILABLE = True
    REPORTLAB_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    REPORTLAB_AVAILABLE = False

from io import BytesIO
from bs4 import BeautifulSoup

from core.models import (
    Carrera, Asignatura, ContenidoAnalitico, FundamentoTeorico, Procedimientos, 
    CalculosResultados, Cuestionario, PracticaLaboratorio, Competencias, 
    ObjetivoPractica, MaterialesHerramientasEquipos, Bibliografia, Titulo
)
from .models import GuiaGenerada
from .forms import GuiaLaboratorioForm, GuiaFilterForm


# ============================================================
# FUNCIÓN HELPER PARA CONVERTIR HTML DE CKEDITOR A WORD
# ============================================================

def limpiar_html_para_word(html_content):
    """
    Limpia el HTML de CKEditor para evitar problemas al convertir a Word.
    Elimina scripts, estilos problemáticos y normaliza el contenido.
    Mantiene imágenes base64 para que funcionen en Word.
    """
    if not html_content or not html_content.strip():
        return ""
    
    from bs4 import BeautifulSoup
    import re
    
    try:
        # Parsear el HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Eliminar scripts y estilos que pueden causar problemas
        for tag in soup.find_all(['script', 'style']):
            tag.decompose()
        
        # Procesar imágenes para asegurar compatibilidad con Word
        for img in soup.find_all('img'):
            # Mantener atributos esenciales
            attrs_to_keep = {}
            
            if img.get('src'):
                src = img['src']
                attrs_to_keep['src'] = src
                
                # Para imágenes base64, asegurar que estén correctamente formateadas
                if src.startswith('data:image'):
                    img_size_mb = len(src) / (1024 * 1024)
                    
                    # Limitar tamaño a 10MB (antes era 50MB, muy grande para Word)
                    if len(src) > 10000000:  # 10MB
                        img.replace_with(soup.new_string(f'[Imagen muy grande ({img_size_mb:.1f} MB) - omitida]'))
                        continue
            
            if img.get('alt'):
                attrs_to_keep['alt'] = img['alt']
            
            # Manejar dimensiones de imagen
            if img.get('width'):
                attrs_to_keep['width'] = img['width']
            if img.get('height'):
                attrs_to_keep['height'] = img['height']
            
            # Si no tiene dimensiones, agregar un ancho por defecto razonable
            if 'width' not in attrs_to_keep and 'height' not in attrs_to_keep:
                attrs_to_keep['width'] = '400'  # Ancho por defecto en píxeles
            
            img.attrs = attrs_to_keep
        
        # Convertir de vuelta a string HTML limpio
        clean_html = str(soup)
        
        # Eliminar entidades HTML problemáticas
        clean_html = clean_html.replace('&nbsp;', ' ')
        
        return clean_html
        
    except Exception as e:
        # Si falla la limpieza, devolver el HTML original
        return html_content


def agregar_html_a_documento(doc, html_content, parser=None):
    """
    Convierte contenido HTML de CKEditor a formato Word y lo agrega al documento.
    Maneja correctamente imágenes base64, negritas, cursivas, tablas, etc.
    
    Args:
        doc: Documento de python-docx
        html_content: String con HTML generado por CKEditor
        parser: Instancia de HtmlToDocx (opcional, se crea si no se proporciona)
    
    Returns:
        El documento actualizado
    """
    if not html_content or not html_content.strip():
        return doc
    
    try:
        # Limpiar el HTML antes de convertir
        html_limpio = limpiar_html_para_word(html_content)
        
        if not html_limpio or not html_limpio.strip():
            return doc
        
        # Procesamiento especial para imágenes base64
        from bs4 import BeautifulSoup
        import base64
        import io
        from docx.shared import Inches
        import re
        
        soup = BeautifulSoup(html_limpio, 'html.parser')
        imagenes_encontradas = soup.find_all('img')
        
        # Convertir imágenes base64 a placeholders únicos
        image_map = {}
        contador_imagenes = 0
        
        for img in soup.find_all('img'):
            src = img.get('src', '')
            
            if src.startswith('data:image'):
                try:
                    # Extraer datos base64
                    if ',' in src:
                        header, encoded = src.split(',', 1)
                    else:
                        continue
                    
                    image_data = base64.b64decode(encoded)
                    
                    # Verificar tamaño de imagen
                    if len(image_data) > 10 * 1024 * 1024:  # 10 MB
                        continue
                    
                    # Crear stream de imagen
                    image_stream = io.BytesIO(image_data)
                    
                    # Determinar ancho
                    width_str = img.get('width', img.get('style', ''))
                    try:
                        # Intentar extraer ancho del atributo width o style
                        if 'width:' in width_str:
                            match = re.search(r'width:\s*(\d+)', width_str)
                            if match:
                                width_pixels = int(match.group(1))
                            else:
                                width_pixels = 400
                        elif width_str.endswith('px'):
                            width_pixels = int(width_str[:-2])
                        elif width_str.isdigit():
                            width_pixels = int(width_str)
                        else:
                            width_pixels = 400
                        
                        # Convertir píxeles a pulgadas (96 DPI)
                        width_inches = width_pixels / 96.0
                        
                        # Limitar ancho máximo a 6 pulgadas
                        if width_inches > 6:
                            width_inches = 6
                        if width_inches < 1:
                            width_inches = 4
                    except:
                        width_inches = 4  # Valor por defecto
                    
                    # Guardar datos de imagen con placeholder único
                    contador_imagenes += 1
                    placeholder = f"___IMAGE_PLACEHOLDER_{contador_imagenes}___"
                    image_map[placeholder] = {
                        'data': image_stream,
                        'width': width_inches
                    }
                    
                    # Reemplazar img con placeholder en un párrafo
                    new_tag = soup.new_tag('p')
                    new_tag.string = placeholder
                    img.replace_with(new_tag)
                    
                except Exception as e:
                    continue
        
        # Convertir el HTML con placeholders usando htmldocx
        html_con_placeholders = str(soup)
        
        # Crear parser si no se proporciona
        if parser is None:
            parser = HtmlToDocx()
        
        # Agregar el HTML al documento (con placeholders)
        parser.add_html_to_document(html_con_placeholders, doc)
        
        # Reemplazar placeholders con imágenes reales
        for paragraph in doc.paragraphs:
            for placeholder, img_data in image_map.items():
                if placeholder in paragraph.text:
                    # Limpiar el texto del placeholder
                    paragraph.text = ''
                    # Agregar la imagen en su lugar
                    run = paragraph.add_run()
                    run.add_picture(img_data['data'], width=Inches(img_data['width']))
        
    except Exception as e:
        # Si falla la conversión HTML, agregar como texto plano
        from bs4 import BeautifulSoup
        try:
            soup = BeautifulSoup(html_content, 'html.parser')
            texto_plano = soup.get_text(separator='\n')
            if texto_plano.strip():
                doc.add_paragraph(texto_plano)
        except:
            # Último recurso: agregar el HTML como está (truncado)
            doc.add_paragraph(str(html_content)[:500] + '...')
    
    return doc


@login_required
def lista_guias(request):
    """Vista para listar todas las guías de laboratorio"""
    
    # Filtros
    filter_form = GuiaFilterForm(request.GET)
    guias = GuiaGenerada.objects.all().select_related('carrera', 'asignatura', 'usuario_creador')
    
    # Aplicar filtros
    if filter_form.is_valid():
        if filter_form.cleaned_data['carrera']:
            guias = guias.filter(carrera=filter_form.cleaned_data['carrera'])
        if filter_form.cleaned_data['semestre']:
            guias = guias.filter(semestre=filter_form.cleaned_data['semestre'])
    
    # Búsqueda por texto
    buscar = request.GET.get('buscar')
    if buscar:
        guias = guias.filter(
            Q(titulo__icontains=buscar) |
            Q(asignatura__nombre__icontains=buscar) |
            Q(contenido_analitico__icontains=buscar)
        )
    
    # Ordenar por fecha de creación descendente
    guias = guias.order_by('-created_at')
    
    # Estadísticas
    total_guias = GuiaGenerada.objects.count()
    guias_con_word = GuiaGenerada.objects.exclude(archivo_word='').count()
    guias_con_pdf = GuiaGenerada.objects.exclude(archivo_pdf='').count()
    
    # Guías de este mes
    from datetime import datetime, timedelta
    inicio_mes = datetime.now().replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    guias_este_mes = GuiaGenerada.objects.filter(created_at__gte=inicio_mes).count()
    
    # Paginación
    from django.core.paginator import Paginator
    paginator = Paginator(guias, 12)  # 12 guías por página
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    # Opciones para filtros
    carreras = Carrera.objects.all()
    semestres = list(range(1, 11))  # Semestres 1-10
    
    # Datos adicionales para estadísticas
    guias_recientes = GuiaGenerada.objects.filter(created_at__gte=inicio_mes)
    
    context = {
        'guias': page_obj,  # Usar el objeto paginado
        'page_obj': page_obj,
        'filter_form': filter_form,
        'carreras': carreras,
        'semestres': semestres,
        'total_guias': total_guias,
        'guias_con_word': guias_con_word,
        'guias_con_pdf': guias_con_pdf,
        'guias_este_mes': guias_este_mes,
        'guias_recientes': guias_recientes,
    }
    
    return render(request, 'guias/lista.html', context)


@login_required
def nueva_guia(request):
    """Vista para crear una nueva guía de laboratorio"""
    
    if request.method == 'POST':
        form = GuiaLaboratorioForm(request.POST)
        if form.is_valid():
            try:
                guia = form.save(commit=False)
                guia.usuario_creador = request.user
                guia.save()
                
                # Generar documentos
                word_file = None
                pdf_file = None
                
                if DOCX_AVAILABLE:
                    word_file = generar_documento_word(guia)
                else:
                    messages.warning(request, 'No se pudo generar el archivo Word. Librería python-docx no disponible.')
                
                if REPORTLAB_AVAILABLE:
                    pdf_file = generar_documento_pdf(guia)
                else:
                    messages.warning(request, 'No se pudo generar el archivo PDF. Librería reportlab no disponible.')
                
                # Guardar archivos
                if word_file:
                    from django.core.files.base import ContentFile
                    # Convertir BytesIO a ContentFile para Django
                    guia.archivo_word.save(
                        f'guia_{guia.id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx',
                        ContentFile(word_file.getvalue()),
                        save=True
                    )
                
                if pdf_file:
                    from django.core.files.base import ContentFile
                    # Convertir BytesIO a ContentFile para Django
                    guia.archivo_pdf.save(
                        f'guia_{guia.id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf',
                        ContentFile(pdf_file.getvalue()),
                        save=True
                    )
                
                messages.success(request, 'Guía creada correctamente.')
                return redirect('guias:detalle', guia_id=guia.id)
                
            except Exception as e:
                messages.error(request, f'Error al crear la guía: {str(e)}')
        else:
            # Mostrar errores específicos
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f'{field}: {error}')
    else:
        form = GuiaLaboratorioForm()
    
    context = {
        'form': form,
        'title': 'Nueva Guía de Laboratorio'
    }
    
    return render(request, 'guias/nueva.html', context)


@login_required
def editar_guia(request, guia_id):
    """Vista para editar una guía de laboratorio existente"""
    
    guia = get_object_or_404(GuiaGenerada, id=guia_id)
    
    # Verificar permisos: solo el creador o admin puede editar
    if not request.user.is_staff and guia.usuario_creador != request.user:
        messages.error(request, 'No tiene permisos para editar esta guía.')
        return redirect('guias:lista')
    
    if request.method == 'POST':
        form = GuiaLaboratorioForm(request.POST, instance=guia)
        if form.is_valid():
            try:
                guia = form.save(commit=False)
                # Mantener el usuario creador original
                guia.save()
                form.save_m2m()  # Guardar relaciones ManyToMany
                
                # Regenerar documentos Word y PDF
                word_file = None
                pdf_file = None
                
                if DOCX_AVAILABLE:
                    word_file = generar_documento_word(guia)
                else:
                    messages.warning(request, 'No se pudo regenerar el archivo Word.')
                
                if REPORTLAB_AVAILABLE:
                    pdf_file = generar_documento_pdf(guia)
                else:
                    messages.warning(request, 'No se pudo regenerar el archivo PDF.')
                
                # Actualizar archivos
                if word_file:
                    from django.core.files.base import ContentFile
                    guia.archivo_word.save(
                        f'guia_{guia.id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx',
                        ContentFile(word_file.getvalue()),
                        save=True
                    )
                
                if pdf_file:
                    from django.core.files.base import ContentFile
                    guia.archivo_pdf.save(
                        f'guia_{guia.id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf',
                        ContentFile(pdf_file.getvalue()),
                        save=True
                    )
                
                messages.success(request, 'Guía actualizada correctamente.')
                return redirect('guias:lista')
                
            except Exception as e:
                messages.error(request, f'Error al actualizar la guía: {str(e)}')
        else:
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f'{field}: {error}')
    else:
        form = GuiaLaboratorioForm(instance=guia)
    
    context = {
        'form': form,
        'guia': guia,
        'title': f'Editar Guía: {guia.titulo}',
        'is_edit': True
    }
    
    return render(request, 'guias/nueva.html', context)


@login_required
def detalle_guia(request, guia_id):
    """Vista que redirige a la lista de guías"""
    return redirect('guias:lista')


@login_required
def eliminar_guia(request, guia_id):
    """Vista para eliminar una guía"""
    
    guia = get_object_or_404(GuiaGenerada, id=guia_id)
    
    if request.method == 'POST':
        titulo = guia.titulo
        
        # Eliminar archivos físicos si existen
        if guia.archivo_word and os.path.exists(guia.archivo_word.path):
            os.remove(guia.archivo_word.path)
        if guia.archivo_pdf and os.path.exists(guia.archivo_pdf.path):
            os.remove(guia.archivo_pdf.path)
            
        guia.delete()
        messages.success(request, f'Guía "{titulo}" eliminada exitosamente.')
        return redirect('guias:lista')
    
    return render(request, 'guias/eliminar.html', {'guia': guia})


@login_required
def descargar_word(request, guia_id):
    """Vista para descargar el archivo Word de una guía"""
    
    guia = get_object_or_404(GuiaGenerada, id=guia_id)
    
    if not guia.archivo_word:
        raise Http404("El archivo Word no está disponible")
    
    try:
        # Limpiar el título para el nombre del archivo (quitar caracteres especiales)
        import re
        titulo_limpio = re.sub(r'[^\w\s-]', '', guia.titulo)
        titulo_limpio = re.sub(r'[-\s]+', '_', titulo_limpio)
        filename = f"Guia_{titulo_limpio[:30]}.docx"
        
        # Leer el archivo
        with open(guia.archivo_word.path, 'rb') as f:
            file_content = f.read()
        
        # Crear respuesta con headers correctos
        response = HttpResponse(
            file_content,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        # Agregar headers de descarga
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        response['Content-Length'] = len(file_content)
        response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        response['Pragma'] = 'no-cache'
        response['Expires'] = '0'
        
        return response
        
    except FileNotFoundError:
        raise Http404("El archivo no se encontró en el servidor")
    except Exception as e:
        raise Http404(f"Error al procesar el archivo: {str(e)}")


@login_required
def descargar_pdf(request, guia_id):
    """Vista para descargar el archivo PDF de una guía"""
    
    guia = get_object_or_404(GuiaGenerada, id=guia_id)
    
    if not guia.archivo_pdf:
        raise Http404("El archivo PDF no está disponible")
    
    try:
        with open(guia.archivo_pdf.path, 'rb') as f:
            response = HttpResponse(f.read(), content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="guia_{guia.id}_{guia.titulo[:30]}.pdf"'
            return response
    except FileNotFoundError:
        raise Http404("El archivo no se encontró en el servidor")


# APIs para dropdowns dinámicos

def api_asignaturas(request):
    """API para obtener asignaturas filtradas por carrera"""
    
    carrera_id = request.GET.get('carrera')
    
    if carrera_id:
        asignaturas = Asignatura.objects.filter(carrera_id=carrera_id).order_by('semestre', 'nombre')
    else:
        asignaturas = Asignatura.objects.all().order_by('semestre', 'nombre')
    
    data = {
        'asignaturas': [
            {
                'id': asignatura.id, 
                'nombre': asignatura.get_nombre_display(), 
                'semestre': asignatura.semestre
            }
            for asignatura in asignaturas
        ]
    }
    
    return JsonResponse(data)


# Funciones auxiliares para generación de documentos

def generar_documento_word(guia):
    """Genera un documento Word con el formato oficial de la EMI"""
    if not DOCX_AVAILABLE:
        return None
        
    try:
        doc = Document()
        
        # Configurar márgenes del documento
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # ========== PÁGINA DE CARÁTULA ==========
        
        # Logo y encabezado de la institución (simulado con texto)
        header_table = doc.add_table(rows=1, cols=1)
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        header_cell = header_table.cell(0, 0)
        header_para = header_cell.paragraphs[0]
        
        # Título EMI
        run_emi = header_para.add_run('EMI')
        run_emi.font.name = 'Arial'
        run_emi.font.size = Pt(36)
        run_emi.font.bold = True
        run_emi.font.color.rgb = RGBColor(41, 84, 144)  # Azul EMI
        
        header_para.add_run('\n')
        
        # Subtítulo
        run_subtitle = header_para.add_run('ESCUELA MILITAR DE INGENIERÍA')
        run_subtitle.font.name = 'Arial'
        run_subtitle.font.size = Pt(14)
        run_subtitle.font.bold = True
        run_subtitle.font.color.rgb = RGBColor(41, 84, 144)
        
        header_para.add_run('\n')
        
        # Nombre del Mariscal
        run_name = header_para.add_run('Mcal. Antonio José de Sucre')
        run_name.font.name = 'Arial'
        run_name.font.size = Pt(12)
        run_name.font.color.rgb = RGBColor(41, 84, 144)
        
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Espacios
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Título principal de la guía
        titulo_principal = doc.add_paragraph('GUÍA DE LABORATORIO')
        titulo_principal.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in titulo_principal.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(24)
            run.font.bold = True
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Nombre de la asignatura
        asignatura_para = doc.add_paragraph(f'NOMBRE DE LA ASIGNATURA: {guia.asignatura.nombre.upper()}')
        asignatura_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in asignatura_para.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(16)
            run.font.bold = True
        
        # Línea decorativa
        line_para = doc.add_paragraph('_' * 50)
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in line_para.runs:
            run.font.color.rgb = RGBColor(255, 193, 7)  # Amarillo EMI
            run.font.bold = True
        
        doc.add_paragraph()
        
        # Salto de página
        doc.add_page_break()
        
        # ========== PÁGINA DE CONTENIDO ==========
        
        # Encabezado con logo y título
        crear_encabezado_pagina(doc, "GUÍA DE LABORATORIO")
        
        doc.add_paragraph()
        
        # Título GUÍA DE LABORATORIO
        titulo_guia = doc.add_heading('GUÍA DE LABORATORIO', 1)
        titulo_guia.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Título CONTENIDO
        contenido_title = doc.add_heading('CONTENIDO', 2)
        contenido_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Tabla de contenido
        contenido_table = doc.add_table(rows=11, cols=2)
        contenido_table.style = 'Table Grid'
        contenido_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Encabezados de la tabla
        contenido_table.cell(0, 0).text = 'PRÁCTICA'
        contenido_table.cell(0, 1).text = 'PÁGINA'
        
        # Formatear encabezados
        for i in range(2):
            cell = contenido_table.cell(0, i)
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Llenar contenido
        for i in range(1, 11):
            contenido_table.cell(i, 0).text = f'PL {i}.'
            contenido_table.cell(i, 1).text = 'Pág. 1'
            
            # Centrar contenido
            for j in range(2):
                contenido_table.cell(i, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Salto de página
        doc.add_page_break()
        
        # ========== PÁGINA PRINCIPAL ==========
        
        # Encabezado con logo y título
        crear_encabezado_pagina(doc, "GUÍA DE LABORATORIO")
        
        # Sección 1: DATOS GENERALES
        datos_title = doc.add_heading('1. DATOS GENERALES', 2)
        datos_title.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        # Crear tabla de datos generales
        datos_table = doc.add_table(rows=9, cols=2)
        datos_table.style = 'Table Grid'
        
        # Datos para la tabla
        datos_info = [
            ('CARRERA:', guia.carrera.get_nombre_display()),
            ('SEMESTRE:', str(guia.semestre)),
            ('ASIGNATURA:', guia.asignatura.nombre),
            ('CONTENIDO ANALÍTICO:', guia.contenido_analitico),
            ('UNIDAD DIDÁCTICA:', guia.unidad_didactica),
            ('DOCENTE:', f'{guia.usuario_creador.first_name} {guia.usuario_creador.last_name}'),
            ('Correo Institucional:', ''),
            ('BIBLIOGRAFÍA DE REFERENCIA:', guia.referencia_bibliografica if guia.referencia_bibliografica else '-'),
            ('', '')  # Fila en blanco
        ]
        
        # Llenar datos
        for i, (label, value) in enumerate(datos_info):
            cell_label = datos_table.cell(i, 0)
            cell_value = datos_table.cell(i, 1)
            
            cell_label.text = label
            cell_value.text = value
            
            # Formatear etiquetas en negrita
            if label:
                cell_label.paragraphs[0].runs[0].font.bold = True
        
        # Fila especial para práctica de laboratorio
        practica_row = datos_table.add_row()
        merged_cell = practica_row.cells[0].merge(practica_row.cells[1])
        merged_cell.text = f'PRÁCTICA DE LABORATORIO N°{guia.numero_practica}: {guia.titulo}'
        merged_cell.paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Secciones principales con contenido real
        # 2. COMPETENCIAS
        seccion_heading = doc.add_heading('2. COMPETENCIAS', 2)
        seccion_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        if guia.competencias:
            doc.add_paragraph(guia.competencias)
        else:
            doc.add_paragraph()
        
        # 3. CRITERIOS DE DESEMPEÑO
        seccion_heading = doc.add_heading('3. CRITERIOS DE DESEMPEÑO', 2)
        seccion_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        if guia.criterios_evaluacion:
            doc.add_paragraph(guia.criterios_evaluacion)
        else:
            doc.add_paragraph()
        
        # 4. OBJETIVO DE LA PRÁCTICA
        seccion_heading = doc.add_heading('4. OBJETIVO DE LA PRÁCTICA DE LABORATORIO', 2)
        seccion_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        if guia.objetivo_general:
            doc.add_paragraph(guia.objetivo_general)
        else:
            doc.add_paragraph()
        
        # 5. MATERIALES, HERRAMIENTAS Y EQUIPOS
        seccion_heading = doc.add_heading('5. MATERIALES, HERRAMIENTAS Y EQUIPOS', 2)
        seccion_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        crear_tabla_materiales(doc)
        
        doc.add_page_break()
        
        # ========== PÁGINA DE CONTINUACIÓN ==========
        crear_encabezado_pagina(doc, "GUÍA DE LABORATORIO")
        
        # Continuación de materiales
        reactivos_title = doc.add_paragraph('DETALLE REACTIVOS:')
        reactivos_title.runs[0].font.bold = True
        
        reactivos_table = doc.add_table(rows=5, cols=2)
        reactivos_table.style = 'Table Grid'
        
        # Encabezado
        reactivos_table.cell(0, 0).text = ''
        reactivos_table.cell(0, 1).text = 'CANTIDAD'
        reactivos_table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
        reactivos_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Herramientas
        doc.add_paragraph()
        herramientas_title = doc.add_paragraph('DETALLE HERRAMIENTAS')
        herramientas_title.runs[0].font.bold = True
        
        herramientas_table = doc.add_table(rows=8, cols=2)
        herramientas_table.style = 'Table Grid'
        
        # Encabezado
        herramientas_table.cell(0, 0).text = ''
        herramientas_table.cell(0, 1).text = 'CANTIDAD'
        herramientas_table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
        herramientas_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # ========== SECCIONES CON CONTENIDO HTML DE CKEDITOR ==========
        
        try:
            # Crear parser HTML una sola vez para reutilizar
            html_parser = HtmlToDocx()
        except Exception as parser_error:
            html_parser = None
        
        # Obtener todos los procedimientos relacionados al contenido analítico de la guía
        if hasattr(guia, 'contenido_analitico_obj'):
            contenido_obj = guia.contenido_analitico_obj
        else:
            # Intentar obtener el contenido analítico por nombre
            try:
                # Buscar por nombre del contenido analítico
                contenido_obj = ContenidoAnalitico.objects.filter(
                    nombre=guia.contenido_analitico
                ).first()
                
                # Si no se encuentra por nombre, intentar por descripción
                if not contenido_obj:
                    contenido_obj = ContenidoAnalitico.objects.filter(
                        descripcion__icontains=guia.contenido_analitico
                    ).first()
                    
            except Exception as search_error:
                contenido_obj = None
        
        # 6. FUNDAMENTO TEÓRICO
        doc.add_paragraph()
        fundamento_heading = doc.add_heading('6. FUNDAMENTO TEÓRICO', 2)
        fundamento_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        if contenido_obj:
            # Obtener y agregar FUNDAMENTOS TEÓRICOS
            fundamentos = FundamentoTeorico.objects.filter(contenido_analitico=contenido_obj).order_by('orden')
            if fundamentos.exists():
                for fundamento in fundamentos:
                    # Contenido HTML del fundamento (sin título adicional)
                    if fundamento.contenido:
                        agregar_html_a_documento(doc, fundamento.contenido, html_parser)
                    # Referencias si existen
                    if fundamento.referencias:
                        doc.add_paragraph()
                        ref_para = doc.add_paragraph()
                        ref_run = ref_para.add_run('Referencias:')
                        ref_run.bold = True
                        agregar_html_a_documento(doc, fundamento.referencias, html_parser)
            else:
                doc.add_paragraph()
        else:
            doc.add_paragraph()
        
        # 7. PROCEDIMIENTO
        doc.add_paragraph()
        procedimiento_heading = doc.add_heading('7. PROCEDIMIENTO', 2)
        procedimiento_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        if contenido_obj:
            # Obtener y agregar PROCEDIMIENTOS
            procedimientos = Procedimientos.objects.filter(contenido_analitico=contenido_obj).order_by('orden', 'numero_paso')
            if procedimientos.exists():
                for proc in procedimientos:
                    # Paso numerado
                    paso_titulo = doc.add_paragraph()
                    paso_run = paso_titulo.add_run(f"Paso {proc.numero_paso}: {proc.titulo_paso}")
                    paso_run.bold = True
                    paso_run.font.size = Pt(12)
                    
                    # Descripción HTML del procedimiento
                    if proc.descripcion:
                        agregar_html_a_documento(doc, proc.descripcion, html_parser)
                    
                    # Tiempo estimado si existe
                    if proc.tiempo_estimado:
                        doc.add_paragraph()
                        tiempo_para = doc.add_paragraph()
                        tiempo_run = tiempo_para.add_run(f"⏱ Tiempo estimado: {proc.tiempo_estimado}")
                        tiempo_run.italic = True
                    
                    # Precauciones si existen
                    if proc.precauciones:
                        doc.add_paragraph()
                        prec_para = doc.add_paragraph()
                        prec_run = prec_para.add_run('⚠️ Precauciones:')
                        prec_run.bold = True
                        agregar_html_a_documento(doc, proc.precauciones, html_parser)
                    
                    # Observaciones si existen
                    if proc.observaciones:
                        doc.add_paragraph()
                        obs_para = doc.add_paragraph()
                        obs_run = obs_para.add_run('📝 Observaciones:')
                        obs_run.bold = True
                        agregar_html_a_documento(doc, proc.observaciones, html_parser)
            else:
                doc.add_paragraph()
        else:
            doc.add_paragraph()
        
        # 8. CÁLCULOS Y RESULTADOS
        doc.add_paragraph()
        calculos_heading = doc.add_heading('8. CÁLCULOS Y RESULTADOS', 2)
        calculos_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        if contenido_obj:
            calculos = CalculosResultados.objects.filter(contenido_analitico=contenido_obj).order_by('orden')
            if calculos.exists():
                for calculo in calculos:
                    # Fórmula si existe
                    if calculo.formula:
                        formula_para = doc.add_paragraph()
                        formula_run = formula_para.add_run(f"📐 Fórmula: {calculo.formula}")
                        formula_run.font.name = 'Courier New'
                        formula_run.bold = True
                    
                    # Procedimiento de cálculo (HTML)
                    if calculo.procedimiento_calculo:
                        agregar_html_a_documento(doc, calculo.procedimiento_calculo, html_parser)
                    
                    # Resultado esperado
                    if calculo.resultado_esperado:
                        doc.add_paragraph()
                        resultado_para = doc.add_paragraph()
                        resultado_run = resultado_para.add_run(f"✓ Resultado esperado: {calculo.resultado_esperado}")
                        resultado_run.bold = True
                    
                    # Unidades
                    if calculo.unidades:
                        unidades_para = doc.add_paragraph(f"Unidades: {calculo.unidades}")
                    
                    # Margen de error
                    if calculo.margen_error:
                        margen_para = doc.add_paragraph(f"±Margen de error: {calculo.margen_error}")
                        margen_para.runs[0].italic = True
            else:
                doc.add_paragraph()
        else:
            doc.add_paragraph()
        
        # 9. CUESTIONARIO
        doc.add_paragraph()
        cuestionario_heading = doc.add_heading('9. CUESTIONARIO', 2)
        cuestionario_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        if contenido_obj:
            # Obtener y agregar CUESTIONARIO
            preguntas = Cuestionario.objects.filter(contenido_analitico=contenido_obj).order_by('orden', 'numero_pregunta')
            if preguntas.exists():
                for pregunta in preguntas:
                    # Número y pregunta
                    pregunta_para = doc.add_paragraph()
                    numero_run = pregunta_para.add_run(f"{pregunta.numero_pregunta}. ")
                    numero_run.bold = True
                    
                    # Pregunta HTML
                    if pregunta.pregunta:
                        agregar_html_a_documento(doc, pregunta.pregunta, html_parser)
                    
                    # Tipo de pregunta
                    doc.add_paragraph()
                    tipo_para = doc.add_paragraph(f"[{pregunta.get_tipo_pregunta_display()}]")
                    tipo_para.runs[0].italic = True
                    tipo_para.runs[0].font.size = Pt(9)
                    
                    # Respuesta esperada (si existe - para el docente)
                    if pregunta.respuesta_esperada:
                        doc.add_paragraph('💡 Respuesta esperada (guía docente):', style='Heading 5')
                        agregar_html_a_documento(doc, pregunta.respuesta_esperada, html_parser)
                    
                    # Puntuación
                    if pregunta.puntuacion:
                        puntos_para = doc.add_paragraph(f"Puntuación: {pregunta.puntuacion} pts")
                        puntos_para.runs[0].font.size = Pt(9)
                    
                    doc.add_paragraph()  # Espaciado entre preguntas
            else:
                doc.add_paragraph("No se han definido preguntas de cuestionario para esta práctica.")
        else:
            doc.add_paragraph()
            doc.add_paragraph()
        
        # Firma al final
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        
        firma_nombre = doc.add_paragraph(f'{guia.usuario_creador.first_name} {guia.usuario_creador.last_name}'.upper())
        firma_nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER
        firma_nombre.runs[0].font.bold = True
        
        firma_cargo = doc.add_paragraph(f'DOCENTE DE LABORATORIO DE LA ASIGNATURA ({guia.asignatura.nombre.upper()})')
        firma_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # ========== GUARDAR DOCUMENTO CON VALIDACIÓN ==========
        try:
            # Guardar en memoria
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # Validar que el buffer tiene contenido
            buffer_size = buffer.getbuffer().nbytes
            if buffer_size == 0:
                print("❌ Error: El documento generado está vacío")
                return None
            
            print(f"✅ Documento Word generado exitosamente: {buffer_size} bytes")
            return buffer
            
        except Exception as save_error:
            print(f"❌ Error al guardar documento Word: {save_error}")
            import traceback
            traceback.print_exc()
            return None
        
    except Exception as e:
        print(f"❌ Error general generando documento Word: {e}")
        import traceback
        traceback.print_exc()
        return None


def crear_encabezado_pagina(doc, titulo_documento):
    """Crea el encabezado estándar de la EMI para páginas internas"""
    # Tabla del encabezado
    header_table = doc.add_table(rows=1, cols=3)
    header_table.style = 'Table Grid'
    
    # Logo EMI (celda izquierda)
    logo_cell = header_table.cell(0, 0)
    logo_para = logo_cell.paragraphs[0]
    logo_run = logo_para.add_run('EMI')
    logo_run.font.name = 'Arial'
    logo_run.font.size = Pt(14)
    logo_run.font.bold = True
    logo_run.font.color.rgb = RGBColor(41, 84, 144)
    
    logo_para.add_run('\n')
    subtitle_run = logo_para.add_run('ESCUELA MILITAR DE INGENIERÍA')
    subtitle_run.font.name = 'Arial'
    subtitle_run.font.size = Pt(8)
    subtitle_run.font.color.rgb = RGBColor(41, 84, 144)
    
    logo_para.add_run('\n')
    name_run = logo_para.add_run('Mcal. Antonio José de Sucre')
    name_run.font.name = 'Arial'
    name_run.font.size = Pt(7)
    name_run.font.color.rgb = RGBColor(41, 84, 144)
    
    # Título del documento (celda central)
    title_cell = header_table.cell(0, 1)
    title_para = title_cell.paragraphs[0]
    title_para.text = titulo_documento
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(12)
    title_run.font.bold = True
    
    # Código y versión (celda derecha)
    code_cell = header_table.cell(0, 2)
    code_para = code_cell.paragraphs[0]
    code_para.text = 'Código:\nVersión: 1'
    code_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in code_para.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(9)
    
    # Ajustar ancho de columnas
    header_table.columns[0].width = Inches(2.5)
    header_table.columns[1].width = Inches(3.5)
    header_table.columns[2].width = Inches(1.5)


def crear_tabla_materiales(doc):
    """Crea las tablas de materiales con el formato específico"""
    
    # Tabla de EQUIPOS
    equipos_title = doc.add_paragraph('DETALLE EQUIPOS:')
    equipos_title.runs[0].font.bold = True
    
    equipos_table = doc.add_table(rows=5, cols=2)
    equipos_table.style = 'Table Grid'
    
    # Encabezado
    equipos_table.cell(0, 0).text = ''
    equipos_table.cell(0, 1).text = 'CANTIDAD'
    equipos_table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
    equipos_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Tabla de MATERIALES
    materiales_title = doc.add_paragraph('DETALLE MATERIALES:')
    materiales_title.runs[0].font.bold = True
    
    materiales_table = doc.add_table(rows=5, cols=2)
    materiales_table.style = 'Table Grid'
    
    # Encabezado
    materiales_table.cell(0, 0).text = ''
    materiales_table.cell(0, 1).text = 'CANTIDAD'
    materiales_table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
    materiales_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def generar_documento_pdf(guia):
    """Genera un documento PDF con el formato oficial de la EMI"""
    if not REPORTLAB_AVAILABLE:
        return None
        
    try:
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter,
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=72)
        
        # Estilos personalizados
        styles = getSampleStyleSheet()
        
        # Estilo para el título EMI
        emi_title_style = ParagraphStyle(
            'EMITitle',
            parent=styles['Normal'],
            fontSize=36,
            textColor=colors.Color(41/255, 84/255, 144/255),  # Azul EMI
            alignment=1,  # Centrado
            fontName='Helvetica-Bold',
            spaceAfter=10
        )
        
        # Estilo para subtítulo EMI
        emi_subtitle_style = ParagraphStyle(
            'EMISubtitle',
            parent=styles['Normal'],
            fontSize=14,
            textColor=colors.Color(41/255, 84/255, 144/255),
            alignment=1,
            fontName='Helvetica-Bold',
            spaceAfter=5
        )
        
        # Estilo para nombre del mariscal
        mariscal_style = ParagraphStyle(
            'MariscalStyle',
            parent=styles['Normal'],
            fontSize=12,
            textColor=colors.Color(41/255, 84/255, 144/255),
            alignment=1,
            fontName='Helvetica',
            spaceAfter=20
        )
        
        # Estilo para título principal
        main_title_style = ParagraphStyle(
            'MainTitle',
            parent=styles['Heading1'],
            fontSize=24,
            alignment=1,
            fontName='Helvetica-Bold',
            spaceAfter=20
        )
        
        # Estilo para nombre de asignatura
        subject_style = ParagraphStyle(
            'SubjectStyle',
            parent=styles['Normal'],
            fontSize=16,
            alignment=1,
            fontName='Helvetica-Bold',
            spaceAfter=10
        )
        
        # Contenido del documento
        story = []
        
        # ========== PÁGINA DE CARÁTULA ==========
        
        # Logo y título EMI
        story.append(Paragraph("EMI", emi_title_style))
        story.append(Paragraph("ESCUELA MILITAR DE INGENIERÍA", emi_subtitle_style))
        story.append(Paragraph("Mcal. Antonio José de Sucre", mariscal_style))
        
        # Espacios
        story.append(Spacer(1, 50))
        
        # Título principal
        story.append(Paragraph("GUÍA DE LABORATORIO", main_title_style))
        story.append(Spacer(1, 30))
        
        # Nombre de asignatura
        story.append(Paragraph(f"NOMBRE DE LA ASIGNATURA: {guia.asignatura.nombre.upper()}", subject_style))
        
        # Línea decorativa
        story.append(Spacer(1, 10))
        story.append(Paragraph("_" * 50, ParagraphStyle('Line', alignment=1, textColor=colors.Color(255/255, 193/255, 7/255))))
        
        story.append(PageBreak())
        
        # ========== PÁGINA DE CONTENIDO ==========
        
        # Encabezado
        crear_encabezado_pdf(story, "GUÍA DE LABORATORIO")
        
        story.append(Paragraph("GUÍA DE LABORATORIO", styles['Heading1']))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("CONTENIDO", styles['Heading2']))
        story.append(Spacer(1, 12))
        
        # Tabla de contenido
        contenido_data = [['PRÁCTICA', 'PÁGINA']]
        for i in range(1, 11):
            contenido_data.append([f'PL {i}.', 'Pág. 1'])
        
        contenido_table = Table(contenido_data, colWidths=[4*inch, 2*inch])
        contenido_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(contenido_table)
        story.append(PageBreak())
        
        # ========== PÁGINA PRINCIPAL ==========
        
        # Encabezado
        crear_encabezado_pdf(story, "GUÍA DE LABORATORIO")
        
        # Datos generales
        story.append(Paragraph("1. DATOS GENERALES", styles['Heading2']))
        
        datos_data = [
            ['CARRERA:', guia.carrera.get_nombre_display()],
            ['SEMESTRE:', str(guia.semestre)],
            ['ASIGNATURA:', guia.asignatura.nombre],
            ['CONTENIDO ANALÍTICO:', guia.contenido_analitico],
            ['UNIDAD DIDÁCTICA:', guia.unidad_didactica],
            ['DOCENTE:', f'{guia.usuario_creador.first_name} {guia.usuario_creador.last_name}'],
            ['Correo Institucional:', ''],
            ['BIBLIOGRAFÍA DE REFERENCIA:', '-\n-\n-\n-']
        ]
        
        datos_table = Table(datos_data, colWidths=[2.5*inch, 4*inch])
        datos_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.white),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        
        story.append(datos_table)
        story.append(Spacer(1, 12))
        
        # Práctica de laboratorio
        practica_data = [[f'PRÁCTICA DE LABORATORIO N°: ...... TÍTULO: {guia.titulo}']]
        practica_table = Table(practica_data, colWidths=[6.5*inch])
        practica_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.white),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        
        story.append(practica_table)
        story.append(Spacer(1, 20))
        
        # Secciones principales
        secciones = [
            '2. COMPETENCIAS',
            '3. CRITERIOS DE DESEMPEÑO',
            '4. OBJETIVO DE LA PRÁCTICA DE LABORATORIO',
            '5. MATERIALES, HERRAMIENTAS Y EQUIPOS'
        ]
        
        for seccion in secciones:
            story.append(Paragraph(seccion, styles['Heading2']))
            
            if seccion == '5. MATERIALES, HERRAMIENTAS Y EQUIPOS':
                # Tablas de materiales
                crear_tablas_materiales_pdf(story)
            else:
                story.append(Spacer(1, 30))
        
        story.append(PageBreak())
        
        # ========== PÁGINA DE CONTINUACIÓN ==========
        crear_encabezado_pdf(story, "GUÍA DE LABORATORIO")
        
        # Continuación de materiales y secciones finales
        story.append(Paragraph("DETALLE REACTIVOS:", ParagraphStyle('Bold', fontName='Helvetica-Bold')))
        
        reactivos_data = [['', 'CANTIDAD']] + [['', ''] for _ in range(4)]
        reactivos_table = Table(reactivos_data, colWidths=[4*inch, 2*inch])
        reactivos_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        story.append(reactivos_table)
        story.append(Spacer(1, 15))
        
        story.append(Paragraph("DETALLE HERRAMIENTAS", ParagraphStyle('Bold', fontName='Helvetica-Bold')))
        
        herramientas_data = [['', 'CANTIDAD']] + [['', ''] for _ in range(7)]
        herramientas_table = Table(herramientas_data, colWidths=[4*inch, 2*inch])
        herramientas_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        story.append(herramientas_table)
        story.append(Spacer(1, 20))
        
        # Secciones finales
        story.append(Paragraph("6. PROCEDIMIENTO", styles['Heading2']))
        story.append(Spacer(1, 40))
        
        story.append(Paragraph("7. CUESTIONARIO", styles['Heading2']))
        story.append(Spacer(1, 60))
        
        # Firma
        story.append(Paragraph(f'{guia.usuario_creador.first_name} {guia.usuario_creador.last_name}'.upper(),
                              ParagraphStyle('signature', alignment=1, fontSize=12, fontName='Helvetica-Bold')))
        story.append(Paragraph(f'DOCENTE DE LABORATORIO DE LA ASIGNATURA ({guia.asignatura.nombre.upper()})',
                              ParagraphStyle('cargo', alignment=1, fontSize=10)))
        
        # Construir PDF
        doc.build(story)
        buffer.seek(0)
        
        return buffer
        
    except Exception as e:
        print(f"Error generando documento PDF: {e}")
        return None


def crear_encabezado_pdf(story, titulo_documento):
    """Crea el encabezado estándar EMI para PDF"""
    header_data = [
        [
            Paragraph("EMI<br/>ESCUELA MILITAR DE INGENIERÍA<br/>Mcal. Antonio José de Sucre", 
                     ParagraphStyle('HeaderLogo', fontSize=8, textColor=colors.Color(41/255, 84/255, 144/255))),
            Paragraph(titulo_documento, 
                     ParagraphStyle('HeaderTitle', fontSize=12, fontName='Helvetica-Bold', alignment=1)),
            Paragraph("Código:<br/>Versión: 1", 
                     ParagraphStyle('HeaderCode', fontSize=9))
        ]
    ]
    
    header_table = Table(header_data, colWidths=[2.5*inch, 3.5*inch, 1.5*inch])
    header_table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ALIGN', (1, 0), (1, 0), 'CENTER'),
    ]))
    
    story.append(header_table)
    story.append(Spacer(1, 20))


def crear_tablas_materiales_pdf(story):
    """Crea las tablas de materiales para PDF"""
    
    # Tabla de equipos
    story.append(Paragraph("DETALLE EQUIPOS:", ParagraphStyle('Bold', fontName='Helvetica-Bold')))
    
    equipos_data = [['', 'CANTIDAD']] + [['', ''] for _ in range(4)]
    equipos_table = Table(equipos_data, colWidths=[4*inch, 2*inch])
    equipos_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    story.append(equipos_table)
    story.append(Spacer(1, 15))
    
    # Tabla de materiales
    story.append(Paragraph("DETALLE MATERIALES:", ParagraphStyle('Bold', fontName='Helvetica-Bold')))
    
    materiales_data = [['', 'CANTIDAD']] + [['', ''] for _ in range(4)]
    materiales_table = Table(materiales_data, colWidths=[4*inch, 2*inch])
    materiales_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    story.append(materiales_table)
    story.append(Spacer(1, 15))


# === NUEVAS VISTAS PARA FUNCIONALIDAD AVANZADA DE GUÍAS ===

@login_required
def generar_guia_pdf_completa(request, guia_id):
    """Genera PDF completo con todos los datos de la guía de laboratorio"""
    
    if not REPORTLAB_AVAILABLE:
        return JsonResponse({'error': 'ReportLab no está disponible'}, status=500)
    
    guia = get_object_or_404(GuiaGenerada, id=guia_id)
    
    # Crear el PDF
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
    
    # Estilos
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=20,
        alignment=1,  # Centrado
        textColor=colors.darkblue
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=12,
        spaceAfter=12,
        textColor=colors.darkgreen
    )
    
    normal_style = styles['Normal']
    
    story = []
    
    # === TÍTULO PRINCIPAL ===
    story.append(Paragraph("GUÍA DE LABORATORIO", title_style))
    story.append(Paragraph(f"<b>{guia.titulo}</b>", title_style))
    story.append(Spacer(1, 20))
    
    # === INFORMACIÓN INSTITUCIONAL ===
    story.append(Paragraph("DATOS INSTITUCIONALES", heading_style))
    
    # Información básica en tabla
    datos_institucionales = [
        ['Unidad Académica:', guia.carrera.unidad_academica.get_nombre_display()],
        ['Carrera:', guia.carrera.get_nombre_display()],
        ['Asignatura:', guia.asignatura.get_nombre_display()],
        ['Semestre:', f"{guia.semestre}° Semestre"],
        ['Tipo de Práctica:', guia.get_tipo_practica_display()],
        ['Duración:', f"{guia.duracion_horas} horas"],
        ['Número de Práctica:', str(guia.numero_practica)],
    ]
    
    datos_table = Table(datos_institucionales, colWidths=[2*inch, 4*inch])
    datos_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
    ]))
    story.append(datos_table)
    story.append(Spacer(1, 20))
    
    # === DATOS DE LA ASIGNATURA ===
    story.append(Paragraph("DATOS DE LA ASIGNATURA", heading_style))
    
    # Información curricular
    if hasattr(guia, 'codigo_competencia') and guia.codigo_competencia:
        story.append(Paragraph(f"<b>Código de Competencia:</b> {guia.codigo_competencia}", normal_style))
    if hasattr(guia, 'sigla_curricular') and guia.sigla_curricular:
        story.append(Paragraph(f"<b>Sigla Curricular:</b> {guia.sigla_curricular}", normal_style))
    if hasattr(guia, 'carga_horaria_semestral') and guia.carga_horaria_semestral:
        story.append(Paragraph(f"<b>Carga Horaria Semestral:</b> {guia.carga_horaria_semestral} horas", normal_style))
    if hasattr(guia, 'carga_horaria_semanal') and guia.carga_horaria_semanal:
        story.append(Paragraph(f"<b>Carga Horaria Semanal:</b> {guia.carga_horaria_semanal} horas", normal_style))
    
    story.append(Spacer(1, 15))
    
    # === UNIDAD DIDÁCTICA ===
    story.append(Paragraph("UNIDAD DIDÁCTICA", heading_style))
    story.append(Paragraph(guia.unidad_didactica or "No especificada", normal_style))
    story.append(Spacer(1, 15))
    
    # === CONTENIDO ANALÍTICO ===
    story.append(Paragraph("CONTENIDO ANALÍTICO", heading_style))
    story.append(Paragraph(guia.contenido_analitico or "No especificado", normal_style))
    story.append(Spacer(1, 15))
    
    # === BIBLIOGRAFÍA ===
    if guia.referencia_bibliografica:
        story.append(Paragraph("BIBLIOGRAFÍA", heading_style))
        story.append(Paragraph(guia.referencia_bibliografica, normal_style))
        story.append(Spacer(1, 15))
    
    # === PRÁCTICA DE LABORATORIO ===
    story.append(Paragraph("PRÁCTICA DE LABORATORIO", heading_style))
    
    # === COMPETENCIAS ===
    story.append(Paragraph("COMPETENCIAS", ParagraphStyle('SubHeading', parent=heading_style, fontSize=11)))
    story.append(Paragraph(guia.competencias or "No especificadas", normal_style))
    story.append(Spacer(1, 10))
    
    # === OBJETIVO ===
    story.append(Paragraph("OBJETIVO DE LA PRÁCTICA", ParagraphStyle('SubHeading', parent=heading_style, fontSize=11)))
    story.append(Paragraph("<b>Objetivo General:</b>", normal_style))
    story.append(Paragraph(guia.objetivo_general or "No especificado", normal_style))
    
    if guia.objetivos_especificos:
        story.append(Paragraph("<b>Objetivos Específicos:</b>", normal_style))
        story.append(Paragraph(guia.objetivos_especificos, normal_style))
    story.append(Spacer(1, 10))
    
    # === FUNDAMENTO TEÓRICO ===
    if hasattr(guia, 'fundamento_teorico') and guia.fundamento_teorico:
        story.append(Paragraph("FUNDAMENTO TEÓRICO", ParagraphStyle('SubHeading', parent=heading_style, fontSize=11)))
        story.append(Paragraph(guia.fundamento_teorico, normal_style))
        story.append(Spacer(1, 10))
    
    # === PROCEDIMIENTOS ===
    story.append(Paragraph("PROCEDIMIENTOS", ParagraphStyle('SubHeading', parent=heading_style, fontSize=11)))
    
    if guia.preparacion_previa:
        story.append(Paragraph("<b>Preparación Previa:</b>", normal_style))
        story.append(Paragraph(guia.preparacion_previa, normal_style))
        story.append(Spacer(1, 10))
    
    story.append(Paragraph("<b>Procedimiento de la Práctica:</b>", normal_style))
    story.append(Paragraph(guia.procedimientos or "No especificado", normal_style))
    story.append(Spacer(1, 15))
    
    # === EQUIPOS ===
    story.append(Paragraph("EQUIPOS", ParagraphStyle('SubHeading', parent=heading_style, fontSize=11)))
    
    if guia.equipos_requeridos.exists():
        equipos_data = [['Equipo', 'Cantidad', 'Estado']]
        for equipo in guia.equipos_requeridos.all()[:10]:  # Limitar a 10 equipos
            equipos_data.append([
                equipo.equipo_existente or 'Sin nombre',
                str(equipo.numero_unidades) if hasattr(equipo, 'numero_unidades') else '1',
                equipo.get_estado_display() if equipo.estado else 'N/A'
            ])
        
        equipos_table = Table(equipos_data, colWidths=[3*inch, 1*inch, 1.5*inch])
        equipos_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
        ]))
        story.append(equipos_table)
    else:
        story.append(Paragraph("No se han especificado equipos para esta práctica.", normal_style))
    
    story.append(Spacer(1, 15))
    
    # === MATERIALES ===
    story.append(Paragraph("MATERIALES", ParagraphStyle('SubHeading', parent=heading_style, fontSize=11)))
    
    if guia.insumos_requeridos.exists():
        materiales_data = [['Material', 'Cantidad', 'Unidad']]
        for insumo in guia.insumos_requeridos.all()[:10]:  # Limitar a 10 materiales
            materiales_data.append([
                insumo.nombre_elemento or 'Sin nombre',
                str(insumo.cantidad),
                insumo.get_unidad_medida_display() if insumo.unidad_medida else 'unidades'
            ])
        
        materiales_table = Table(materiales_data, colWidths=[3*inch, 1*inch, 1.5*inch])
        materiales_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgreen),
        ]))
        story.append(materiales_table)
    else:
        story.append(Paragraph("No se han especificado materiales para esta práctica.", normal_style))
    
    story.append(Spacer(1, 15))
    
    # === HERRAMIENTAS ===
    story.append(Paragraph("HERRAMIENTAS", ParagraphStyle('SubHeading', parent=heading_style, fontSize=11)))
    
    # Filtrar herramientas de los insumos
    herramientas = guia.insumos_requeridos.filter(categoria='herramientas')
    if herramientas.exists():
        herramientas_data = [['Herramienta', 'Cantidad', 'Estado']]
        for herramienta in herramientas[:10]:
            herramientas_data.append([
                herramienta.nombre_elemento or 'Sin nombre',
                str(herramienta.cantidad),
                herramienta.get_estado_display() if herramienta.estado else 'N/A'
            ])
        
        herramientas_table = Table(herramientas_data, colWidths=[3*inch, 1*inch, 1.5*inch])
        herramientas_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightyellow),
        ]))
        story.append(herramientas_table)
    else:
        story.append(Paragraph("No se han especificado herramientas para esta práctica.", normal_style))
    
    story.append(Spacer(1, 20))
    
    # === CÁLCULOS Y RESULTADOS ===
    story.append(Paragraph("CÁLCULOS Y RESULTADOS", ParagraphStyle('SubHeading', parent=heading_style, fontSize=11)))
    
    # Espacio para completar
    calculo_data = [
        ['Descripción del Cálculo', 'Fórmula', 'Resultado'],
        ['', '', ''],
        ['', '', ''],
        ['', '', ''],
    ]
    
    calculo_table = Table(calculo_data, colWidths=[2*inch, 2*inch, 1.5*inch])
    calculo_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightcoral),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
    ]))
    story.append(calculo_table)
    story.append(Spacer(1, 20))
    
    # === CUESTIONARIO ===
    story.append(Paragraph("CUESTIONARIO", ParagraphStyle('SubHeading', parent=heading_style, fontSize=11)))
    story.append(Paragraph(guia.cuestionario or "No se ha definido cuestionario para esta práctica.", normal_style))
    
    # Generar PDF
    doc.build(story)
    
    # Preparar respuesta
    buffer.seek(0)
    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    filename = f"Guia_Laboratorio_{guia.titulo.replace(' ', '_')}.pdf"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    return response


@login_required 
def detalle_guia_completa(request, guia_id):
    """Vista de detalle completa de la guía con todos los campos"""
    
    guia = get_object_or_404(GuiaGenerada, id=guia_id)
    
    context = {
        'guia': guia,
        'equipos': guia.equipos_requeridos.all()[:20],  # Limitar para rendimiento
        'materiales': guia.insumos_requeridos.filter(categoria__in=['materiales', 'reactivos'])[:20],
        'herramientas': guia.insumos_requeridos.filter(categoria='herramientas')[:20],
        'puede_editar': request.user == guia.usuario_creador or request.user.is_staff,
    }
    
    return render(request, 'guias/detalle_completa.html', context)


# ===== NUEVAS VISTAS PARA PRÁCTICAS DE LABORATORIO =====

def html_to_latex(html_content):
    """Convierte HTML de CKEditor a LaTeX, escapando caracteres especiales"""
    if not html_content:
        return ""
    
    # Extraer texto limpio del HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    text = soup.get_text()
    
    # Escapar caracteres especiales de LaTeX
    replacements = {
        '\\': '\\textbackslash{}',
        '&': '\\&',
        '%': '\\%',
        '$': '\\$',
        '#': '\\#',
        '_': '\\_',
        '{': '\\{',
        '}': '\\}',
        '~': '\\textasciitilde{}',
        '^': '\\textasciicircum{}',
    }
    
    for old, new in replacements.items():
        text = text.replace(old, new)
    
    # Convertir saltos de línea dobles en párrafos LaTeX
    text = text.replace('\n\n', '\n\n\\par ')
    
    return text


@login_required
def generar_practica_word(request, practica_id):
    """Generar PDF con LaTeX de una práctica de laboratorio - VERSIÓN COMPLETA EMI"""
    
    try:
        # Obtener la práctica y modelos relacionados  
        practica = get_object_or_404(PracticaLaboratorio, id=practica_id)
        asignatura = practica.contenido_analitico.unidad_didactica.asignatura
        unidad = practica.contenido_analitico.unidad_didactica
        contenido = practica.contenido_analitico
        
        print(f"📄 Generando PDF LaTeX para práctica: {practica.nombre}")
        
        # Crear directorio temporal
        temp_dir = tempfile.mkdtemp()
        tex_file = os.path.join(temp_dir, 'practica.tex')
        
        # ============================================================
        # GENERAR CONTENIDO LaTeX COMPLETO
        # ============================================================
        
        latex_content = r"""\documentclass[11pt,a4paper]{article}
\usepackage[utf8]{inputenc}
\usepackage[spanish]{babel}
\usepackage{geometry}
\usepackage{graphicx}
\usepackage{xcolor}
\usepackage{fancyhdr}
\usepackage{array}
\usepackage{longtable}
\usepackage{booktabs}
\usepackage{enumitem}
\usepackage{titlesec}
\usepackage{tcolorbox}

% Configuración de página
\geometry{a4paper, left=2.5cm, right=2.5cm, top=2.5cm, bottom=2.5cm}

% Colores EMI
\definecolor{emiazul}{RGB}{0,46,93}
\definecolor{emidorado}{RGB}{255,215,0}
\definecolor{emigris}{RGB}{245,245,245}

% Formato de secciones
\titleformat{\section}
  {\normalfont\Large\bfseries\color{emiazul}}{\thesection}{1em}{}
  [\titlerule]
  
\titleformat{\subsection}
  {\normalfont\large\bfseries\color{emiazul}}{\thesubsection}{1em}{}

% Encabezado y pie de página (no aplicar a la portada)
\fancypagestyle{contenido}{
  \fancyhf{}
  \fancyhead[L]{\small\textbf{EMI - """ + (asignatura.carrera.nombre if (asignatura and asignatura.carrera) else 'N/A') + r"""}}
  \fancyhead[R]{\small\textbf{Práctica """ + str(practica.orden if practica.orden else 1) + r"""}}
  \fancyfoot[C]{\thepage}
}

\begin{document}

% ============================================================
% PORTADA
% ============================================================
\thispagestyle{empty}

% Logo EMI centrado
\begin{center}
\includegraphics[width=14.3cm,height=19.6cm,keepaspectratio]{emi_logo.png}
\end{center}

\vfill

% Título centrado
\begin{center}
{\Huge\textbf{GUÍA DE LABORATORIO}}\\[1.5cm]
{\Large\textbf{""" + html_to_latex(asignatura.nombre.upper()) + r"""}}
\end{center}

\vfill

% Salto de página para comenzar contenido
\newpage

% ============================================================
% PÁGINA 2: ÍNDICE DE CONTENIDOS
% ============================================================
\thispagestyle{empty}

\begin{center}
{\Large\textbf{GUÍA DE LABORATORIO}}
\end{center}

\vspace{0.5cm}

\noindent\makebox[\linewidth]{\rule{\textwidth}{0.4pt}}

\vspace{1cm}

\begin{center}
{\large\textbf{CONTENIDO}}
\end{center}

\vspace{1cm}

\noindent\textbf{1. DATOS GENERALES} \dotfill Pág. \pageref{sec:datos}\\[0.5cm]
\noindent\textbf{2. COMPETENCIAS} \dotfill Pág. \pageref{sec:competencias}\\[0.5cm]
\noindent\textbf{3. CRITERIOS DE DESEMPEÑO} \dotfill Pág. \pageref{sec:criterios}\\[0.5cm]
\noindent\textbf{4. OBJETIVO DE LA PRÁCTICA DE LABORATORIO} \dotfill Pág. \pageref{sec:objetivo}\\[0.5cm]
\noindent\textbf{5. FUNDAMENTO TEÓRICO} \dotfill Pág. \pageref{sec:fundamento}\\[0.5cm]
\noindent\textbf{6. MATERIALES, HERRAMIENTAS Y EQUIPOS} \dotfill Pág. \pageref{sec:materiales}\\[0.5cm]
\noindent\textbf{7. PROCEDIMIENTO} \dotfill Pág. \pageref{sec:procedimiento}\\[0.5cm]
\noindent\textbf{8. CÁLCULOS Y RESULTADOS} \dotfill Pág. \pageref{sec:calculos}\\[0.5cm]
\noindent\textbf{9. CUESTIONARIO} \dotfill Pág. \pageref{sec:cuestionario}\\[0.5cm]

\vfill

\newpage

% Activar estilo de encabezado para el contenido
\pagestyle{contenido}
\setcounter{page}{3}

% ============================================================
% 1. DATOS GENERALES
% ============================================================
\section{DATOS GENERALES}\label{sec:datos}

\begin{tabular}{|p{5cm}|p{10cm}|}
\hline
\textbf{CARRERA} & """ + (asignatura.carrera.get_nombre_display() if (asignatura and asignatura.carrera) else 'N/A') + r""" \\
\hline
\textbf{SEMESTRE} & """ + str(asignatura.semestre if (asignatura and asignatura.semestre) else 'N/A') + r""" \\
\hline
\textbf{ASIGNATURA} & """ + (asignatura.nombre.upper().replace('_', ' ') if asignatura else 'N/A') + r""" \\
\hline
"""
        
        # Determinar qué mostrar como "Contenido Analítico"
        # Usar siempre el nombre del contenido analítico (no la unidad ni la práctica)
        contenido_analitico_text = contenido.nombre
        
        latex_content += r"""\textbf{CONTENIDO ANALÍTICO} & """ + html_to_latex(contenido_analitico_text.upper()) + r""" \\
\hline
\textbf{UNIDAD DIDÁCTICA} & """ + html_to_latex(unidad.nombre.upper() if unidad else 'N/A') + r""" \\
\hline
\textbf{DOCENTE} & \rule{8cm}{0.4pt} \\
\hline
\textbf{Correo Institucional} & \rule{8cm}{0.4pt} \\
\hline
"""
        
        # Agregar bibliografía
        bibliografias = Bibliografia.objects.filter(contenido_analitico=contenido).order_by('orden')
        biblio_text = ""
        if bibliografias.exists():
            biblio_items = [html_to_latex(b.titulo) for b in bibliografias[:3]]
            biblio_text = ", ".join(biblio_items)
        else:
            biblio_text = "N/A"
        
        latex_content += r"""\textbf{BIBLIOGRAFÍA DE REFERENCIA} & """ + biblio_text + r""" \\
\hline
\end{tabular}

\vspace{0.5cm}

"""
        
        # Usar el Título del formulario si existe, si no usar el nombre de la práctica
        titulo_practica = practica.nombre
        titulos = Titulo.objects.filter(contenido_analitico=contenido).order_by('orden')
        if titulos.exists():
            titulo_practica = titulos.first().texto
        
        latex_content += r"""\begin{tcolorbox}[colback=emidorado,colframe=emiazul,arc=0mm,boxrule=1pt]
\begin{center}
\large\textbf{PRÁCTICA DE LABORATORIO N° """ + str(practica.orden if practica.orden else 1) + r"""}\\[0.3cm]
\Large\textbf{TÍTULO: """ + html_to_latex(titulo_practica.upper()) + r"""}
\end{center}
\end{tcolorbox}

\vspace{0.5cm}

"""
        
        # ============================================================
        # 2. COMPETENCIAS
        # ============================================================
        latex_content += r"""\section{COMPETENCIAS}\label{sec:competencias}

"""
        
        competencias = Competencias.objects.filter(contenido_analitico=contenido).order_by('orden')
        if competencias.exists():
            latex_content += r"""\begin{longtable}{|p{15cm}|}
\hline
"""
            for comp in competencias:
                latex_content += f"{html_to_latex(comp.descripcion)} \\\\\n\\hline\n"
            latex_content += r"\end{longtable}" + "\n\n"
        else:
            latex_content += "No hay competencias definidas para esta práctica.\n\n"
        
        # ============================================================
        # 3. CRITERIOS DE DESEMPEÑO
        # ============================================================
        latex_content += r"""\section{CRITERIOS DE DESEMPEÑO}\label{sec:criterios}

"""
        
        # Mostrar SOLO el Criterio de Desempeño seleccionado del dropdown (si existe)
        if contenido.criterio_desempeno:
            latex_content += r"""\begin{longtable}{|p{15cm}|}
\hline
"""
            latex_content += html_to_latex(contenido.criterio_desempeno.nombre) + r" \\" + "\n\\hline\n"
            if contenido.criterio_desempeno.descripcion:
                latex_content += html_to_latex(contenido.criterio_desempeno.descripcion) + r" \\" + "\n\\hline\n"
            latex_content += r"\end{longtable}" + "\n\n"
        else:
            latex_content += "No hay criterios de desempeño definidos.\n\n"
        
        # ============================================================
        # 4. OBJETIVO DE LA PRÁCTICA DE LABORATORIO
        # ============================================================
        latex_content += r"""\section{OBJETIVO DE LA PRÁCTICA DE LABORATORIO}\label{sec:objetivo}

"""
        
        # Mostrar TODOS los objetivos (incluyendo desempeño) sin tipo ni numeración
        objetivos = ObjetivoPractica.objects.filter(
            contenido_analitico=contenido
        ).order_by('orden')
        
        if objetivos.exists():
            latex_content += r"""\begin{longtable}{|p{15cm}|}
\hline
"""
            for obj in objetivos:
                latex_content += html_to_latex(obj.descripcion) + r" \\" + "\n\\hline\n"
            latex_content += r"\end{longtable}" + "\n\n"
        else:
            latex_content += "No hay objetivos definidos.\n\n"
        
        # ============================================================
        # 5. FUNDAMENTO TEÓRICO
        # ============================================================
        latex_content += r"""\section{FUNDAMENTO TEÓRICO}\label{sec:fundamento}

"""
        
        fundamentos = FundamentoTeorico.objects.filter(contenido_analitico=contenido).order_by('orden')
        
        if fundamentos.exists():
            latex_content += r"""\begin{longtable}{|p{15cm}|}
\hline
"""
            for fund in fundamentos:
                latex_content += html_to_latex(fund.contenido) + r" \\" + "\n\\hline\n"
            latex_content += r"\end{longtable}" + "\n\n"
        else:
            latex_content += "No hay fundamento teórico definido.\n\n"
        
        # ============================================================
        # 6. MATERIALES, HERRAMIENTAS Y EQUIPOS
        # ============================================================
        latex_content += r"""\section{MATERIALES, HERRAMIENTAS Y EQUIPOS}\label{sec:materiales}

"""
        
        # Obtener todos los elementos agrupados por tipo
        equipos = MaterialesHerramientasEquipos.objects.filter(
            contenido_analitico=contenido, 
            tipo_elemento='equipo'
        ).order_by('orden')
        
        materiales = MaterialesHerramientasEquipos.objects.filter(
            contenido_analitico=contenido, 
            tipo_elemento='material'
        ).order_by('orden')
        
        herramientas = MaterialesHerramientasEquipos.objects.filter(
            contenido_analitico=contenido, 
            tipo_elemento='herramienta'
        ).order_by('orden')
        
        reactivos = MaterialesHerramientasEquipos.objects.filter(
            contenido_analitico=contenido, 
            tipo_elemento='reactivo'
        ).order_by('orden')
        
        # 6.1 Equipos
        if equipos.exists():
            latex_content += r"""\subsection{Equipos}

\begin{longtable}{|p{10cm}|p{3cm}|}
\hline
\textbf{Equipo} & \textbf{Cantidad} \\
\hline
"""
            for equipo in equipos:
                latex_content += f"{html_to_latex(equipo.nombre)} & {equipo.cantidad if equipo.cantidad else '1'} \\\\\n\\hline\n"
            
            latex_content += r"""\end{longtable}

"""
        
        # 6.2 Materiales
        if materiales.exists():
            latex_content += r"""\subsection{Materiales}

\begin{longtable}{|p{10cm}|p{3cm}|}
\hline
\textbf{Material} & \textbf{Cantidad} \\
\hline
"""
            for material in materiales:
                latex_content += f"{html_to_latex(material.nombre)} & {material.cantidad if material.cantidad else '1'} \\\\\n\\hline\n"
            
            latex_content += r"""\end{longtable}

"""
        
        # 6.3 Herramientas
        if herramientas.exists():
            latex_content += r"""\subsection{Herramientas}

\begin{longtable}{|p{10cm}|p{3cm}|}
\hline
\textbf{Herramienta} & \textbf{Cantidad} \\
\hline
"""
            for herramienta in herramientas:
                latex_content += f"{html_to_latex(herramienta.nombre)} & {herramienta.cantidad if herramienta.cantidad else '1'} \\\\\n\\hline\n"
            
            latex_content += r"""\end{longtable}

"""
        
        # 6.4 Reactivos
        if reactivos.exists():
            latex_content += r"""\subsection{Reactivos}

\begin{longtable}{|p{10cm}|p{3cm}|}
\hline
\textbf{Reactivo} & \textbf{Cantidad} \\
\hline
"""
            for reactivo in reactivos:
                latex_content += f"{html_to_latex(reactivo.nombre)} & {reactivo.cantidad if reactivo.cantidad else '1'} \\\\\n\\hline\n"
            
            latex_content += r"""\end{longtable}

"""
        
        if not (equipos.exists() or materiales.exists() or herramientas.exists() or reactivos.exists()):
            latex_content += "No hay materiales, equipos ni herramientas definidos.\n\n"
        
        # ============================================================
        # 7. PROCEDIMIENTO
        # ============================================================
        latex_content += r"""\section{PROCEDIMIENTO}\label{sec:procedimiento}

"""
        
        procedimientos = Procedimientos.objects.filter(contenido_analitico=contenido).order_by('numero_paso', 'orden')
        
        if procedimientos.exists():
            latex_content += r"""\begin{longtable}{|p{15cm}|}
\hline
"""
            for proc in procedimientos:
                latex_content += html_to_latex(proc.descripcion) + r" \\" + "\n\\hline\n"
            latex_content += r"\end{longtable}" + "\n\n"
        else:
            latex_content += "No hay procedimiento definido.\n\n"
        
        # ============================================================
        # 8. CÁLCULOS Y RESULTADOS
        # ============================================================
        latex_content += r"""\section{CÁLCULOS Y RESULTADOS}\label{sec:calculos}

"""
        
        calculos = CalculosResultados.objects.filter(contenido_analitico=contenido).order_by('orden')
        
        if calculos.exists():
            latex_content += r"""\begin{longtable}{|p{15cm}|}
\hline
"""
            for calc in calculos:
                if calc.formula:
                    latex_content += f"\\textbf{{Fórmula:}} {html_to_latex(calc.formula)}\n\n"
                
                if calc.procedimiento_calculo:
                    latex_content += html_to_latex(calc.procedimiento_calculo) + r" \\" + "\n\\hline\n"
            latex_content += r"\end{longtable}" + "\n\n"
        else:
            latex_content += "No hay cálculos y resultados definidos.\n\n"
        
        # ============================================================
        # 9. CUESTIONARIO
        # ============================================================
        latex_content += r"""\section{CUESTIONARIO}\label{sec:cuestionario}

"""
        
        cuestionario = Cuestionario.objects.filter(contenido_analitico=contenido).order_by('numero_pregunta', 'orden')
        
        if cuestionario.exists():
            latex_content += r"""\begin{longtable}{|p{15cm}|}
\hline
"""
            for pregunta in cuestionario:
                latex_content += html_to_latex(pregunta.pregunta) + r" \\" + "\n\\hline\n"
            latex_content += r"\end{longtable}" + "\n\n"
        else:
            latex_content += "No hay cuestionario definido.\n\n"
        
        # ============================================================
        # FIRMA DEL DOCENTE
        # ============================================================
        latex_content += r"""
\vspace{3cm}

\begin{center}
\rule{10cm}{0.4pt}\\[0.3cm]
\textbf{GRADO Y NOMBRE}\\[0.5cm]
\textbf{DOCENTE DE LABORATORIO DE LA ASIGNATURA """ + html_to_latex(asignatura.nombre.upper()) + r"""}
\end{center}

\end{document}"""
        
        # Escribir archivo .tex
        with open(tex_file, 'w', encoding='utf-8') as f:
            f.write(latex_content)
        
        print(f"✅ Archivo LaTeX generado: {tex_file}")
        
        # Copiar el logo de EMI al directorio temporal
        logo_source = os.path.join(settings.BASE_DIR, 'static', 'images', 'emi_logo.png')
        logo_dest = os.path.join(temp_dir, 'emi_logo.png')
        
        if os.path.exists(logo_source):
            shutil.copy2(logo_source, logo_dest)
            print(f"✅ Logo copiado: {logo_dest}")
        else:
            print(f"⚠️ Logo no encontrado en: {logo_source}")
        
        # Compilar con pdflatex (dos veces para resolver referencias)
        for i in range(2):
            result = subprocess.run(
                ['pdflatex', '-interaction=nonstopmode', '-output-directory', temp_dir, tex_file],
                capture_output=True,
                text=True,
                timeout=30
            )
            if i == 0:
                print(f"✅ Primera compilación completada")
        
        pdf_file = os.path.join(temp_dir, 'practica.pdf')
        
        if not os.path.exists(pdf_file):
            print(f"❌ Error: PDF no generado")
            print(f"Log: {result.stdout[-1000:]}")
            shutil.rmtree(temp_dir, ignore_errors=True)
            return JsonResponse({'error': 'No se pudo generar el PDF'}, status=500)
        
        print(f"✅ PDF generado: {pdf_file}")
        
        # Leer el PDF generado
        with open(pdf_file, 'rb') as f:
            pdf_data = f.read()
        
        # Limpiar directorio temporal
        shutil.rmtree(temp_dir, ignore_errors=True)
        
        # Preparar respuesta HTTP
        filename = f"Practica_{practica.orden if practica.orden else 'N'}_{practica.nombre[:30]}.pdf"
        filename = filename.replace(' ', '_').replace('/', '_')
        
        response = HttpResponse(pdf_data, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return JsonResponse({'error': str(e)}, status=500)


@login_required
def generar_practica_pdf(request, practica_id):
    """Generar PDF de una práctica de laboratorio (guía basada en práctica real)"""
    
    if not REPORTLAB_AVAILABLE:
        return JsonResponse({'error': 'ReportLab no está disponible'}, status=500)
    
    from core.models import PracticaLaboratorio
    
    # Obtener la práctica
    practica = get_object_or_404(PracticaLaboratorio, id=practica_id)
    asignatura = practica.contenido_analitico.unidad_didactica.asignatura
    unidad = practica.contenido_analitico.unidad_didactica
    contenido = practica.contenido_analitico
    
    # Crear respuesta HTTP con PDF
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="Guia_{practica.nombre.replace(" ", "_")}.pdf"'
    
    # Crear documento PDF
    doc = SimpleDocTemplate(response, pagesize=letter)
    story = []
    styles = getSampleStyleSheet()
    
    # Estilos personalizados
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=30,
        textColor=colors.darkblue,
        alignment=1  # Centrado
    )
    
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=12,
        textColor=colors.darkgreen
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=6
    )
    
    # ===== CONTENIDO DEL PDF =====
    
    # Título principal
    story.append(Paragraph(f"GUÍA DE LABORATORIO", title_style))
    story.append(Paragraph(f"{practica.nombre.upper()}", title_style))
    story.append(Spacer(1, 20))
    
    # Información académica
    story.append(Paragraph("INFORMACIÓN ACADÉMICA", subtitle_style))
    
    data_academica = [
        ['Carrera:', asignatura.carrera],
        ['Asignatura:', asignatura.nombre],
        ['Semestre:', f"{asignatura.semestre}°"],
        ['Unidad Didáctica:', unidad.nombre],
        ['Duración:', f"{practica.duracion_horas} horas"],
        ['Tipo de Práctica:', practica.get_tipo_practica_display()],
        ['Número de Estudiantes:', str(practica.numero_estudiantes)],
    ]
    
    tabla_academica = Table(data_academica, colWidths=[2*inch, 4*inch])
    tabla_academica.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('BACKGROUND', (1, 0), (1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    story.append(tabla_academica)
    story.append(Spacer(1, 20))
    
    # Competencias
    competencias = contenido.competencias.all()
    if competencias.exists():
        story.append(Paragraph("COMPETENCIAS", subtitle_style))
        for i, competencia in enumerate(competencias, 1):
            story.append(Paragraph(f"{i}. {competencia.descripcion}", normal_style))
        story.append(Spacer(1, 15))
    
    # Objetivos de la práctica
    objetivos = contenido.objetivos_practica.all()
    if objetivos.exists():
        story.append(Paragraph("OBJETIVOS DE LA PRÁCTICA", subtitle_style))
        for i, objetivo in enumerate(objetivos, 1):
            story.append(Paragraph(f"{i}. {objetivo.descripcion}", normal_style))
        story.append(Spacer(1, 15))
    
    # Descripción del contenido analítico
    story.append(Paragraph("DESCRIPCIÓN", subtitle_style))
    story.append(Paragraph(contenido.descripcion, normal_style))
    story.append(Spacer(1, 20))
    
    # Información adicional
    story.append(Paragraph("INFORMACIÓN ADICIONAL", subtitle_style))
    
    info_adicional = [
        ['Fecha de generación:', timezone.now().strftime('%d/%m/%Y %H:%M')],
        ['Generado por:', request.user.get_full_name() or request.user.username],
        ['Sistema:', 'Centralización de Laboratorios - UMSA'],
    ]
    
    tabla_info = Table(info_adicional, colWidths=[2*inch, 4*inch])
    tabla_info.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.lightblue),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    story.append(tabla_info)
    
    # Construir PDF
    doc.build(story)
    
    return response


@login_required
def detalle_practica_completa(request, practica_id):
    """Vista de detalle completa de una práctica de laboratorio"""
    
    from core.models import PracticaLaboratorio
    
    practica = get_object_or_404(PracticaLaboratorio, id=practica_id)
    asignatura = practica.contenido_analitico.unidad_didactica.asignatura
    unidad = practica.contenido_analitico.unidad_didactica
    contenido = practica.contenido_analitico
    
    context = {
        'practica': practica,
        'asignatura': asignatura,
        'unidad': unidad,
        'contenido': contenido,
        'competencias': contenido.competencias.all(),
        'objetivos': contenido.objetivos_practica.all(),
        'puede_editar': request.user.is_staff,  # Solo staff puede editar prácticas del currículo
    }
    
    return render(request, 'guias/detalle_practica_completa.html', context)
