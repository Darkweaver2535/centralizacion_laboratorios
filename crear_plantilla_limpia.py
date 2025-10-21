#!/usr/bin/env python3
"""
Script para crear una plantilla Word limpia para las guías de laboratorio.
Esto evita problemas de parsing de Jinja2 causados por etiquetas XML fragmentadas.
"""

import os
import django
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'centralizacion.settings')
django.setup()

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def crear_plantilla_limpia():
    """Crea una plantilla Word limpia con variables de Jinja2"""
    
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    # Header con logo y datos de la universidad
    header_table = doc.add_table(rows=3, cols=3)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Fila 1 - Logo y título
    row1 = header_table.rows[0]
    row1.cells[0].text = "[LOGO EMI]"
    row1.cells[1].text = "ESCUELA MILITAR DE INGENIERÍA\n\"MCAL. ANTONIO JOSÉ DE SUCRE\"\nBOLIVIA"
    row1.cells[2].text = "{{codigo_de_documento}}"
    
    # Fila 2 - Datos académicos
    row2 = header_table.rows[1]
    row2.cells[0].text = ""
    row2.cells[1].text = "UNIDAD ACADÉMICA: {{unidad_academica}}\nCARRERA: {{carrera}}\nASIGNATURA: {{nombre_de_la_asignatura}}"
    row2.cells[2].text = "Versión: {{version_de_documento}}\nFecha: {{fecha_de_documento}}"
    
    # Fila 3 - Título del documento
    row3 = header_table.rows[2]
    row3.cells[0].text = ""
    row3.cells[1].text = "GUÍA DE LABORATORIO"
    row3.cells[2].text = "Página: {{pagina}}"
    
    doc.add_paragraph()
    
    # Título de la práctica
    titulo = doc.add_heading("{{titulo}}", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Sección 1: Información General
    doc.add_heading("1. INFORMACIÓN GENERAL", level=2)
    
    info_table = doc.add_table(rows=6, cols=2)
    info_data = [
        ["Docente:", "{{docente}}"],
        ["Auxiliar:", "{{auxiliar}}"],
        ["Fecha:", "{{fecha}}"],
        ["Duración:", "{{duracion}} minutos"],
        ["Grupo:", "{{grupo}}"],
        ["Semestre:", "{{semestre}}"]
    ]
    
    for i, (label, value) in enumerate(info_data):
        info_table.cell(i, 0).text = label
        info_table.cell(i, 1).text = value
    
    # Sección 2: Competencias
    doc.add_heading("2. COMPETENCIAS", level=2)
    doc.add_paragraph("{{competencias}}")
    
    # Sección 3: Fundamentación Teórica
    doc.add_heading("3. FUNDAMENTACIÓN TEÓRICA", level=2)
    doc.add_paragraph("{{fundamentacion_teorica}}")
    
    # Sección 4: Objetivos
    doc.add_heading("4. OBJETIVOS", level=2)
    doc.add_heading("4.1. Objetivo General", level=3)
    doc.add_paragraph("{{objetivo_general}}")
    
    doc.add_heading("4.2. Objetivos Específicos", level=3)
    doc.add_paragraph("{{objetivos_especificos}}")
    
    # Sección 5: Equipos y Materiales
    doc.add_heading("5. EQUIPOS Y MATERIALES", level=2)
    
    # Equipos
    doc.add_heading("5.1. Equipos", level=3)
    equipos_table = doc.add_table(rows=6, cols=3)
    equipos_table.style = 'Table Grid'
    
    # Headers de equipos
    hdr_cells = equipos_table.rows[0].cells
    hdr_cells[0].text = 'Ítem'
    hdr_cells[1].text = 'Equipo'
    hdr_cells[2].text = 'Cantidad'
    
    # Filas de equipos
    for i in range(1, 6):
        row = equipos_table.rows[i]
        row.cells[0].text = str(i)
        row.cells[1].text = f"{{{{equipo{i}}}}}"
        row.cells[2].text = f"{{{{cantidad_equipo{i}}}}}"
    
    # Materiales
    doc.add_heading("5.2. Materiales", level=3)
    materiales_table = doc.add_table(rows=6, cols=3)
    materiales_table.style = 'Table Grid'
    
    # Headers de materiales
    hdr_cells = materiales_table.rows[0].cells
    hdr_cells[0].text = 'Ítem'
    hdr_cells[1].text = 'Material'
    hdr_cells[2].text = 'Cantidad'
    
    # Filas de materiales
    for i in range(1, 6):
        row = materiales_table.rows[i]
        row.cells[0].text = str(i)
        row.cells[1].text = f"{{{{material{i}}}}}"
        row.cells[2].text = f"{{{{cantidad_material{i}}}}}"
    
    # Reactivos
    doc.add_heading("5.3. Reactivos", level=3)
    reactivos_table = doc.add_table(rows=6, cols=3)
    reactivos_table.style = 'Table Grid'
    
    # Headers de reactivos
    hdr_cells = reactivos_table.rows[0].cells
    hdr_cells[0].text = 'Ítem'
    hdr_cells[1].text = 'Reactivo'
    hdr_cells[2].text = 'Cantidad'
    
    # Filas de reactivos
    for i in range(1, 6):
        row = reactivos_table.rows[i]
        row.cells[0].text = str(i)
        row.cells[1].text = f"{{{{reactivo{i}}}}}"
        row.cells[2].text = f"{{{{cantidad_reactivo{i}}}}}"
    
    # Herramientas
    doc.add_heading("5.4. Herramientas", level=3)
    herramientas_table = doc.add_table(rows=6, cols=3)
    herramientas_table.style = 'Table Grid'
    
    # Headers de herramientas
    hdr_cells = herramientas_table.rows[0].cells
    hdr_cells[0].text = 'Ítem'
    hdr_cells[1].text = 'Herramienta'
    hdr_cells[2].text = 'Cantidad'
    
    # Filas de herramientas
    for i in range(1, 6):
        row = herramientas_table.rows[i]
        row.cells[0].text = str(i)
        row.cells[1].text = f"{{{{herramienta{i}}}}}"
        row.cells[2].text = f"{{{{cantidad_herramienta{i}}}}}"
    
    # Sección 6: Procedimiento
    doc.add_heading("6. PROCEDIMIENTO", level=2)
    doc.add_paragraph("{{procedimiento}}")
    
    # Sección 7: Parte Experimental
    doc.add_heading("7. PARTE EXPERIMENTAL", level=2)
    doc.add_paragraph("{{parte_experimental}}")
    
    # Sección 8: Criterios de Desempeño
    doc.add_heading("8. CRITERIOS DE DESEMPEÑO", level=2)
    doc.add_paragraph("{{criterios_de_desempeno}}")
    
    # Sección 9: Conclusiones
    doc.add_heading("9. CONCLUSIONES", level=2)
    doc.add_paragraph("{{conclusiones}}")
    
    # Sección 10: Referencias
    doc.add_heading("10. REFERENCIAS BIBLIOGRÁFICAS", level=2)
    doc.add_paragraph("{{referencias_bibliograficas}}")
    
    return doc

def main():
    """Función principal"""
    print("🔄 Creando plantilla Word limpia...")
    
    doc = crear_plantilla_limpia()
    
    # Guardar la plantilla
    output_path = "/Users/alvaroencinas/Desktop/centralizacion_laboratorios/templates/core/plantilla_guia_laboratorio_limpia.docx"
    
    # Crear directorio si no existe
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    doc.save(output_path)
    
    print(f"✅ Plantilla creada en: {output_path}")
    print(f"📋 La plantilla incluye:")
    print(f"   • Header oficial EMI con variables")
    print(f"   • 10 secciones estructuradas")
    print(f"   • Tablas para equipos, materiales, reactivos y herramientas")
    print(f"   • Variables Jinja2 limpias (sin fragmentación XML)")
    print(f"   • Formato profesional con estilos apropiados")
    
    print(f"\n🔧 Variables incluidas en la plantilla:")
    variables = [
        "codigo_de_documento", "version_de_documento", "fecha_de_documento", "pagina",
        "unidad_academica", "carrera", "nombre_de_la_asignatura", "titulo",
        "docente", "auxiliar", "fecha", "duracion", "grupo", "semestre",
        "competencias", "fundamentacion_teorica", "objetivo_general", "objetivos_especificos",
        "procedimiento", "parte_experimental", "criterios_de_desempeno", "conclusiones", 
        "referencias_bibliograficas"
    ]
    
    # Variables de recursos (equipos, materiales, etc.)
    for categoria in ["equipo", "material", "reactivo", "herramienta"]:
        for i in range(1, 6):
            variables.append(f"{categoria}{i}")
            variables.append(f"cantidad_{categoria}{i}")
    
    for i, var in enumerate(variables, 1):
        print(f"   {i:2d}. {{{{{var}}}}}")
    
    print(f"\n📄 Total de variables: {len(variables)}")
    
    return output_path

if __name__ == "__main__":
    main()