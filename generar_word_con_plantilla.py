"""
Script temporal para generar Word usando plantilla
Este código reemplazará la función generar_practica_word en guias/views.py
"""

from django.shortcuts import get_object_or_404
from django.http import HttpResponse, JsonResponse
from django.contrib.auth.decorators import login_required
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from io import BytesIO
import os
import re
from htmldocx import HtmlToDocx

@login_required
def generar_practica_word(request, practica_id):
    """Generar Word de una práctica usando plantilla con marcadores {{ variable }}"""
    
    from core.models import PracticaLaboratorio, FundamentoTeorico, Procedimientos, CalculosResultados, Cuestionario
    from django.conf import settings
    
    try:
        # Obtener la práctica
        practica = get_object_or_404(PracticaLaboratorio, id=practica_id)
        asignatura = practica.contenido_analitico.unidad_didactica.asignatura
        unidad = practica.contenido_analitico.unidad_didactica
        contenido = practica.contenido_analitico
        
        print(f"📄 Generando Word para práctica: {practica.nombre}")
        
        # Cargar plantilla
        plantilla_path = os.path.join(settings.BASE_DIR, 'pruebas', 'FORMATO GUÍA DE LABORATORIO.docx')
        
        if not os.path.exists(plantilla_path):
            return JsonResponse({'error': f'Plantilla no encontrada en {plantilla_path}'}, status=500)
        
        print(f"✅ Cargando plantilla: {plantilla_path}")
        doc = Document(plantilla_path)
        
        # === DATOS PARA REEMPLAZAR ===
        replacements = {
            '{{ nombre_de_la_asignatura }}': asignatura.nombre,
            '{{ parte_indice }}': str(practica.orden if practica.orden else 1),
            '{{ pagina }}': '1',
            '{{ asignatura }}': asignatura.nombre,
            '{{ carrera }}': asignatura.carrera.nombre if asignatura.carrera else 'N/A',
            '{{ semestre }}': str(asignatura.semestre) if asignatura.semestre else 'N/A',
            '{{ nombre_practica }}': practica.nombre,
            '{{ duracion }}': f"{practica.duracion_horas} horas",
            '{{ numero_estudiantes }}': str(practica.numero_estudiantes) if practica.numero_estudiantes else 'Según capacidad',
        }
        
        # Reemplazar en párrafos
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, value)
        
        # Reemplazar en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                for run in paragraph.runs:
                                    if key in run.text:
                                        run.text = run.text.replace(key, value)
        
        # === RELLENAR SECCIONES ESPECÍFICAS ===
        
        # Inicializar parser HTML
        html_parser = None
        try:
            html_parser = HtmlToDocx()
            print("✅ Parser HTML creado")
        except Exception as e:
            print(f"⚠️ Error creando parser HTML: {e}")
        
        # Función auxiliar para agregar contenido HTML
        def agregar_html_a_celda(cell, html_content):
            """Agrega contenido HTML a una celda de tabla"""
            if not html_content:
                return
            
            # Limpiar HTML
            html_limpio = limpiar_html_para_word(html_content)
            
            # Limpiar párrafos existentes en la celda
            for p in cell.paragraphs:
                p.clear()
            
            if html_parser and html_limpio:
                try:
                    html_parser.add_html_to_cell(html_limpio, cell)
                except Exception as e:
                    print(f"⚠️ Error agregando HTML: {e}")
                    # Fallback: agregar como texto plano
                    from bs4 import BeautifulSoup
                    soup = BeautifulSoup(html_limpio, 'html.parser')
                    cell.text = soup.get_text()
            else:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(html_content, 'html.parser')
                cell.text = soup.get_text()
        
        # TABLA 2: COMPETENCIAS
        if len(doc.tables) > 2:
            competencias_cell = doc.tables[2].rows[1].cells[0]
            competencias = contenido.competencias.all() if hasattr(contenido, 'competencias') else []
            if competencias:
                texto_competencias = '\n'.join([f"• {comp.descripcion}" for comp in competencias])
                competencias_cell.text = texto_competencias
        
        # TABLA 3: CRITERIOS DE DESEMPEÑO (puede venir de objetivos)
        if len(doc.tables) > 3:
            criterios_cell = doc.tables[3].rows[1].cells[0]
            objetivos = contenido.objetivos_practica.all() if hasattr(contenido, 'objetivos_practica') else []
            if objetivos:
                texto_objetivos = '\n'.join([f"{i+1}. {obj.descripcion}" for i, obj in enumerate(objetivos)])
                criterios_cell.text = texto_objetivos
        
        # TABLA 4: OBJETIVOS
        if len(doc.tables) > 4:
            objetivos_cell = doc.tables[4].rows[1].cells[0]
            if objetivos:
                texto_objetivos = '\n'.join([f"• {obj.descripcion}" for obj in objetivos])
                objetivos_cell.text = texto_objetivos
        
        # TABLA 5: FUNDAMENTOS TEÓRICOS
        if len(doc.tables) > 5:
            fundamentos_cell = doc.tables[5].rows[1].cells[0]
            fundamentos = FundamentoTeorico.objects.filter(contenido_analitico=contenido).order_by('orden')
            if fundamentos.exists():
                # Limpiar celda
                fundamentos_cell.text = ''
                for fund in fundamentos:
                    # Agregar título
                    p_titulo = fundamentos_cell.add_paragraph()
                    run_titulo = p_titulo.add_run(fund.titulo)
                    run_titulo.bold = True
                    run_titulo.font.size = Pt(12)
                    
                    # Agregar contenido HTML
                    if fund.contenido:
                        agregar_html_a_celda(fundamentos_cell, fund.contenido)
                    
                    # Agregar referencias si existen
                    if fund.referencias:
                        p_ref = fundamentos_cell.add_paragraph()
                        run_ref = p_ref.add_run('Referencias:')
                        run_ref.bold = True
                        agregar_html_a_celda(fundamentos_cell, fund.referencias)
        
        # TABLA 6: MATERIALES, HERRAMIENTAS Y EQUIPOS
        if len(doc.tables) > 6:
            materiales_table = doc.tables[6]
            materiales_equipos = contenido.materiales_herramientas_equipos.all() if hasattr(contenido, 'materiales_herramientas_equipos') else []
            
            # Limpiar filas existentes (excepto encabezado)
            for _ in range(len(materiales_table.rows) - 1):
                materiales_table._element.remove(materiales_table.rows[-1]._element)
            
            # Agregar datos
            if materiales_equipos:
                for item in materiales_equipos:
                    row = materiales_table.add_row()
                    row.cells[0].text = item.nombre
                    row.cells[1].text = item.cantidad or '1'
        
        # TABLA 7: PROCEDIMIENTO
        if len(doc.tables) > 7:
            procedimiento_cell = doc.tables[7].rows[1].cells[0]
            procedimientos = Procedimientos.objects.filter(contenido_analitico=contenido).order_by('orden', 'numero_paso')
            
            if procedimientos.exists():
                procedimiento_cell.text = ''
                for proc in procedimientos:
                    # Título del paso
                    p_paso = procedimiento_cell.add_paragraph()
                    run_paso = p_paso.add_run(f"Paso {proc.numero_paso}: {proc.titulo_paso}")
                    run_paso.bold = True
                    run_paso.font.size = Pt(11)
                    
                    # Descripción con HTML
                    if proc.descripcion:
                        agregar_html_a_celda(procedimiento_cell, proc.descripcion)
                    
                    # Tiempo estimado
                    if proc.tiempo_estimado:
                        p_tiempo = procedimiento_cell.add_paragraph()
                        run_tiempo = p_tiempo.add_run(f"⏱ Tiempo estimado: {proc.tiempo_estimado}")
                        run_tiempo.italic = True
                    
                    # Precauciones
                    if proc.precauciones:
                        p_prec = procedimiento_cell.add_paragraph()
                        run_prec = p_prec.add_run('⚠️ Precauciones:')
                        run_prec.bold = True
                        agregar_html_a_celda(procedimiento_cell, proc.precauciones)
        
        # TABLA 8: CÁLCULOS Y RESULTADOS
        if len(doc.tables) > 8:
            calculos_cell = doc.tables[8].rows[1].cells[0]
            calculos = CalculosResultados.objects.filter(contenido_analitico=contenido).order_by('orden')
            
            if calculos.exists():
                calculos_cell.text = ''
                for calc in calculos:
                    # Título
                    p_titulo = calculos_cell.add_paragraph()
                    run_titulo = p_titulo.add_run(calc.titulo)
                    run_titulo.bold = True
                    run_titulo.font.size = Pt(11)
                    
                    # Fórmula
                    if calc.formula:
                        p_formula = calculos_cell.add_paragraph()
                        run_formula = p_formula.add_run(f"📐 Fórmula: {calc.formula}")
                        run_formula.font.name = 'Courier New'
                    
                    # Procedimiento de cálculo con HTML
                    if calc.procedimiento_calculo:
                        agregar_html_a_celda(calculos_cell, calc.procedimiento_calculo)
                    
                    # Resultado esperado
                    if calc.resultado_esperado:
                        p_resultado = calculos_cell.add_paragraph()
                        run_resultado = p_resultado.add_run(f"✓ Resultado esperado: {calc.resultado_esperado}")
                        run_resultado.bold = True
        
        # TABLA 9: CUESTIONARIO
        if len(doc.tables) > 9:
            cuestionario_cell = doc.tables[9].rows[1].cells[0]
            cuestionarios = Cuestionario.objects.filter(contenido_analitico=contenido).order_by('orden', 'numero_pregunta')
            
            if cuestionarios.exists():
                cuestionario_cell.text = ''
                for preg in cuestionarios:
                    # Número de pregunta
                    p_num = cuestionario_cell.add_paragraph()
                    run_num = p_num.add_run(f"{preg.numero_pregunta}. ")
                    run_num.bold = True
                    
                    # Pregunta con HTML
                    if preg.pregunta:
                        agregar_html_a_celda(cuestionario_cell, preg.pregunta)
                    
                    # Tipo de pregunta
                    p_tipo = cuestionario_cell.add_paragraph()
                    run_tipo = p_tipo.add_run(f"[{preg.get_tipo_pregunta_display()}]")
                    run_tipo.italic = True
                    run_tipo.font.size = Pt(9)
                    
                    # Respuesta esperada (para guía docente)
                    if preg.respuesta_esperada:
                        p_resp = cuestionario_cell.add_paragraph()
                        run_resp = p_resp.add_run('💡 Respuesta esperada (guía docente):')
                        run_resp.bold = True
                        run_resp.font.size = Pt(9)
                        agregar_html_a_celda(cuestionario_cell, preg.respuesta_esperada)
        
        # === GUARDAR DOCUMENTO ===
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Validar tamaño
        buffer_size = buffer.getbuffer().nbytes
        if buffer_size == 0:
            print("❌ Documento vacío")
            return JsonResponse({'error': 'El documento está vacío'}, status=500)
        
        print(f"✅ Documento generado: {buffer_size} bytes")
        
        # Crear respuesta HTTP
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        # Nombre de archivo
        titulo_limpio = re.sub(r'[^\w\s-]', '', practica.nombre)
        titulo_limpio = re.sub(r'[-\s]+', '_', titulo_limpio)
        filename = f"Guia_Lab_{titulo_limpio[:40]}.docx"
        
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        response['Content-Length'] = buffer_size
        
        return response
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return JsonResponse({'error': str(e)}, status=500)
